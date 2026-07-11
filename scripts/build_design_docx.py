"""Build a formatted Architecture Design Document as .docx."""
from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

OUT = r"C:\Users\vchaurasia\Downloads\Concall_Intelligence_Architecture_Design.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)
GREEN = RGBColor(0x37, 0x56, 0x23)
AMBER = RGBColor(0xBF, 0x8F, 0x00)
GREY = RGBColor(0x44, 0x54, 0x6A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

NAVY_H = "1F3864"
BLUE_H = "2E5496"
GREEN_H = "375623"
AMBER_H = "BF8F00"
GREY_H = "44546A"
LBLUE_H = "D9E1F2"
LBLUE2_H = "EAF0FA"
LGREEN_H = "E2EFDA"
LAMBER_H = "FFF2CC"
LGREY_H = "F2F2F2"

doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK

for m in ("top", "bottom", "left", "right"):
    setattr(doc.sections[0], f"{m}_margin", Inches(0.8))


def shade(cell, hex_):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_)
    tcPr.append(sh)


def set_cell_text(cell, text, *, bold=False, color=None, size=10.5, mono=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Consolas" if mono else "Calibri"
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    return p


def banner(text, fill_hex, size=14):
    """Full-width colored heading paragraph via a 1-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    shade(cell, fill_hex)
    set_cell_text(cell, text, bold=True, color=WHITE, size=size)
    _no_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _set_table_borders(table, color="BFBFBF", sz="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def data_table(headers, rows, widths=None, header_hex=NAVY_H, zebra=True,
               mono_cols=(), bold_cols=(0,)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], header_hex)
        set_cell_text(hdr[i], h, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        band = LGREY_H if (zebra and r_idx % 2 == 1) else "FFFFFF"
        for i, val in enumerate(row):
            shade(cells[i], band)
            set_cell_text(cells[i], val, bold=(i in bold_cols), mono=(i in mono_cols))
    _set_table_borders(t)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def para(text, *, italic=False, size=10.5, color=DARK, space=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space)
    return p


def mono_block(lines, fill_hex=LGREY_H):
    t = doc.add_table(rows=len(lines), cols=1)
    for i, ln in enumerate(lines):
        cell = t.rows[i].cells[0]
        shade(cell, fill_hex)
        set_cell_text(cell, ln, mono=True, size=9.5)
    _set_table_borders(t, color="D9D9D9")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ============================ TITLE PAGE ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Concall Intelligence Platform")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Architecture Design Document (ADD)")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = BLUE

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Quarterly earnings-call intelligence & stock-ranking platform. Tracks management "
                "guidance \u2192 delivery, scores companies vs a 10-section Business Analysis Template, "
                "computes PEG valuation, and ranks stocks.\nDecision-support only \u2014 not investment advice.")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = GREY
doc.add_paragraph()

data_table(
    ["Field", "Value"],
    [
        ["Document version", "1.0"],
        ["Date", "2026-07-11"],
        ["Owner", "@vchaurasia_microsoft"],
        ["Project root", r"C:\Users\vchaurasia\concall-intel\\"],
        ["Runtime", "Python 3 \u00b7 FastAPI \u00b7 SQLAlchemy \u00b7 SQLite \u00b7 uvicorn 127.0.0.1:8010"],
        ["Status", "Implemented \u2014 working prototype, 23 companies live"],
    ],
    widths=[1.9, 4.9], header_hex=GREY_H, bold_cols=(0,),
)

banner("Table of Contents", GREY_H, size=13)
data_table(
    ["Section", "Contents"],
    [
        ["1. Purpose & Scope", "Problem, goals, non-goals"],
        ["2. Architectural Principles", "Principles & how they are realized"],
        ["3. Logical Architecture", "5-layer stack (presentation \u2192 persistence)"],
        ["4. Component Architecture", "Module-by-module responsibilities"],
        ["5. Data Model", "Entities, FK contract, temporal model"],
        ["6. Data Flows", "Ingestion + ranking-refresh flows"],
        ["7. Valuation (PEG)", "Section 9 formula & composite integration"],
        ["8. Scalability & Deployment", "Prototype vs scale-out view"],
        ["9. Decisions, Risks & Future", "Trade-offs, risks/mitigations, roadmap"],
    ],
    widths=[2.6, 4.2], header_hex=BLUE_H, bold_cols=(0,),
)

doc.add_page_break()

# ============================ 1. PURPOSE & SCOPE ============================
banner("1. Purpose & Scope", NAVY_H)
para("Problem statement", size=11.5, color=NAVY, space=2)
para("Investors read dozens of concall transcripts per company across years but cannot easily answer: "
     "did management actually deliver on what they promised? Manual tracking of guidance \u2192 outcome "
     "across 30 companies \u00d7 20 quarters is infeasible.")

banner("Goals", GREEN_H, size=12)
data_table(["#", "Goal"],
    [[str(i), g] for i, g in enumerate([
        "Persist ALL transcripts and extracted facts so history is never lost.",
        "Auto-extract metrics, guidance and red flags from each concall.",
        "Build guidance \u2192 outcome chains across quarters (FY24 \u2192 FY25 \u2192 FY26).",
        "Score each company 0-100 vs a 10-section template (incl. Section 9 PEG valuation).",
        "Rank all companies 1 \u2192 N on a composite of business quality, CAGR, ROE and valuation.",
        "Ingest new quarters incrementally; update only the affected company's dashboard.",
        "Be horizontally scalable and provider-agnostic (LLM, extractor, DB all swappable).",
    ], 1)],
    widths=[0.5, 6.3], header_hex=GREEN_H, bold_cols=())

banner("Non-goals", AMBER_H, size=12)
data_table(["#", "Out of scope"],
    [[str(i), g] for i, g in enumerate([
        "Real-time price feeds / trading execution.",
        "Portfolio management / broker integration.",
        "Guaranteeing external market-data completeness (gaps are flagged, never fabricated).",
    ], 1)],
    widths=[0.5, 6.3], header_hex=AMBER_H, bold_cols=())

doc.add_page_break()

# ============================ 2. PRINCIPLES ============================
banner("2. Architectural Principles", NAVY_H)
data_table(["#", "Principle", "How it is realized"],
    [[str(i), p, h] for i, (p, h) in enumerate([
        ("Separation of concerns", "Ingestion, Extraction, Analysis, Guidance, Scoring, API, UI are independent modules under app/."),
        ("Provider abstraction", "LLM (stub/openai/azure/agent) and extractor (pymupdf/pdfplumber/azure_docintel) selected by config; swap without code changes."),
        ("Idempotent pipeline", "Re-running any transcript is safe; each stage checks TranscriptStatus before acting."),
        ("DB-backed memory", "Every fact is a row; the 'compare against history' feature is a query, not a cache."),
        ("Stateless API", "FastAPI app holds no session state \u2192 trivially replicated behind a load balancer."),
        ("Storage portability", "SQLAlchemy ORM \u2192 SQLite today, Postgres tomorrow, zero API-contract change."),
        ("Honest data", "Missing data flagged (*, N/A, source flags M/C/G/P) rather than guessed."),
    ], 1)],
    widths=[0.4, 2.0, 4.4], header_hex=NAVY_H, bold_cols=(1,))

doc.add_page_break()

# ============================ 3. LOGICAL ARCHITECTURE ============================
banner("3. Logical Architecture \u2014 5-Layer Stack", NAVY_H)
data_table(["Layer", "Name", "Modules", "Responsibility"],
    [
        ["L5", "Presentation", "app/static: index.html, app.js, style.css", "Dashboard/leaderboards, per-company Insights (10 sections + PEG card), ranking table."],
        ["L4", "API", "app/main.py, api/routes_*", "FastAPI REST: upload, insights, /fundamental-ranking, leaderboards. Stateless."],
        ["L3", "Domain / Service", "app/services/*, workers/pipeline.py", "ingestion \u2192 extraction \u2192 analysis \u2192 guidance \u2192 scoring, orchestrated per transcript."],
        ["L2", "Persistence", "app/db.py, app/models.py", "SQLAlchemy ORM over SQLite (concall.db); 10 entity tables."],
        ["L1", "External / Offline", "scripts/*.py, data/*.json", "Yahoo Finance (CAGR, PE), BSE/NSE/Screener transcripts; offline analytics \u2192 DB."],
    ],
    widths=[0.6, 1.4, 2.4, 2.4], header_hex=NAVY_H, bold_cols=(0, 1), mono_cols=(2,))
para("Request/data paths:", size=11, color=BLUE, space=2)
mono_block([
    "PDF upload  -> API -> Service pipeline -> Persistence      (write path)",
    "Browser     -> API -> Persistence -> JSON -> SPA render     (read path)",
    "Scripts     -> Yahoo/BSE -> data/*.json -> persist -> DB    (enrichment path)",
])

doc.add_page_break()

# ============================ 4. COMPONENTS ============================
banner("4. Component Architecture", NAVY_H)
data_table(["#", "Component", "Module", "Responsibility"],
    [[str(i), nm, mod, resp] for i, (nm, mod, resp) in enumerate([
        ("Ingestion", "services/ingestion.py", "Store uploaded PDF blob, create Transcript(status=uploaded), parse period_ordinal (FY*4+Q)."),
        ("Extraction", "services/extraction.py", "PDF bytes \u2192 text (PyMuPDF default; pdfplumber/azure_docintel/stub). status=extracted."),
        ("Analysis", "services/analysis.py + llm/", "LLMClient \u2192 TranscriptAnalysis: summary, sentiment, aggression, metrics/guidance/red_flags/section_scores."),
        ("LLM providers", "llm/client.py", "StubLLM (offline, deterministic), OpenAI/Azure (structured outputs), AgentFileLLM (real analysis, no key)."),
        ("Guidance reconciler", "services/guidance.py", "Match open promises vs new-quarter actuals \u2192 MET/PARTIAL/MISSED/IN_PROGRESS + GuidanceOutcome."),
        ("Scoring", "services/scoring.py", "10 section scores \u2192 total_0_100 + growth/reliability/aggression/consistency + composite."),
        ("Fundamental ranking", "scripts/* + FundamentalRanking", "Business quality + Stock CAGR + Sales CAGR + ROE + PEG \u2192 composite \u2192 Rank 1..N."),
        ("Presentation", "app/static/*", "Vanilla-JS SPA (no build step); renders leaderboards, insights, PEG cards."),
    ], 1)],
    widths=[0.4, 1.6, 2.1, 2.7], header_hex=NAVY_H, bold_cols=(1,), mono_cols=(2,))

doc.add_page_break()

# ============================ 5. DATA MODEL ============================
banner("5. Data Model \u2014 Entities", NAVY_H)
data_table(["#", "Entity", "Keyed by", "Notes"],
    [[str(i), e, k, n] for i, (e, k, n) in enumerate([
        ("Company", "id", "One row per stock (name, ticker)."),
        ("Transcript", "id, company_id", "One row per concall PDF; status uploaded\u2192extracted\u2192reconciled\u2192scored; period_ordinal."),
        ("Metric", "transcript_id", "Extracted financial metric (revenue/ebitda/pat/margin)."),
        ("Guidance", "company_id + source_transcript_id", "Management commitment (NOT transcript_id)."),
        ("GuidanceOutcome", "guidance_id", "Whether guidance was met; the promise\u2192outcome chain node."),
        ("RedFlag", "transcript_id", "Risk/concern detected in commentary."),
        ("SectionScore", "company_id + as_of_period_ordinal", "0-10 per template section."),
        ("CompanyScore", "company_id + as_of_period_ordinal", "Score col = total_0_100; also composite, growth_score, guidance_reliability_score."),
        ("ResearchSheet", "company_id", "Generated research output."),
        ("FundamentalRanking", "company_id", "rank, CAGRs, ROE + 5 PEG columns."),
    ], 1)],
    widths=[0.4, 1.7, 2.2, 2.5], header_hex=NAVY_H, bold_cols=(1,), mono_cols=(2,))

banner("Foreign-key contract (respect on deletes/joins)", AMBER_H, size=12)
data_table(["Table", "Keyed by"],
    [
        ["Metric, RedFlag", "transcript_id"],
        ["Guidance", "company_id + source_transcript_id"],
        ["GuidanceOutcome", "guidance_id"],
        ["SectionScore, CompanyScore", "company_id (+ as_of_period_ordinal)"],
    ],
    widths=[2.6, 4.2], header_hex=AMBER_H, bold_cols=(0,), mono_cols=(1,))
para("Temporal model: period_ordinal = FY*4 + Q gives a total order over quarters, powering guidance "
     "due-dates, as-of scoring snapshots, and multi-year chain assembly without date parsing at query time.",
     italic=True, color=GREY)

doc.add_page_break()

# ============================ 6. DATA FLOWS ============================
banner("6.1 New Transcript Ingestion (incremental)", NAVY_H)
data_table(["Step", "Stage", "Action"],
    [[str(i), s, a] for i, (s, a) in enumerate([
        ("Ingestion", "Store blob + create Transcript(status=uploaded); parse period_ordinal."),
        ("Extraction", "PDF \u2192 text; status=extracted."),
        ("Analysis", "LLM \u2192 metrics / guidance / red_flags / section_scores."),
        ("Guidance", "Reconcile prior open promises vs this quarter's actuals; status=reconciled."),
        ("Scoring", "Recompute CompanyScore as-of this quarter; status=scored."),
        ("Dashboard", "That company's insights update automatically on next API read."),
    ], 1)],
    widths=[0.6, 1.5, 4.7], header_hex=NAVY_H, bold_cols=(1,))

banner("6.2 Ranking Refresh (offline batch)", GREEN_H, size=12)
data_table(["Step", "Script", "Purpose"],
    [[str(i), s, a] for i, (s, a) in enumerate([
        ("build_fundamental_ranking.py", "Recompute composite + PEG factor."),
        ("persist_fundamental_ranking.py", "Write ranking rows to DB."),
        ("fetch_valuation_peg.py", "Refresh trailing PE + growth (if stale)."),
        ("persist_valuation_peg.py", "Write the 5 PEG columns."),
    ], 1)],
    widths=[0.6, 3.0, 3.2], header_hex=GREEN_H, bold_cols=(), mono_cols=(1,))

doc.add_page_break()

# ============================ 7. VALUATION (PEG) ============================
banner("7. Section 9 \u2014 PEG Valuation", NAVY_H)
data_table(["#", "Aspect", "Detail"],
    [[str(i), k, v] for i, (k, v) in enumerate([
        ("Formula", "PEG = trailing_PE / growth_pct"),
        ("Verdict", "PEG < 1 = Undervalued ; else Overvalued ; N/A if PE or growth missing."),
        ("Growth priority", "C = computed Profit CAGR \u2192 S = sustainable Sales CAGR \u2192 Y = Yahoo YoY earnings growth (volatile, flagged)."),
        ("Composite weight", "Valuation factor = 20% of composite; = inverted min-max of PEG capped at 4.0 (low PEG boosts)."),
        ("Missing data", "No PEG \u2192 no boost, never a penalty. Loss-makers / no-growth = N/A (GOKEX, HFCL, RAIN, STLTECH, IDEAFORGE)."),
        ("Observed effect", "SANSERA #11\u2192#4, SAMBHV #17\u2192#8 (both PEG<1); ACUTAAS / NEOGEN dropped."),
    ], 1)],
    widths=[0.4, 1.8, 4.6], header_hex=NAVY_H, bold_cols=(1,))

banner("Composite scoring weights", GREEN_H, size=12)
data_table(["Factor", "Weight", "Source"],
    [
        ["Business Quality", "30%", "10-section template score"],
        ["Stock Price CAGR", "25%", "Yahoo chart API (winsorized 100%)"],
        ["Sales CAGR", "15%", "Transcript-mined (flags M/C/G)"],
        ["ROE", "10%", "Transcript-mined (P = ROCE-proxy)"],
        ["Valuation (PEG)", "20%", "PEG = PE / growth; low PEG \u2192 boost"],
    ],
    widths=[2.0, 1.0, 3.8], header_hex=GREEN_H, bold_cols=(0,))

doc.add_page_break()

# ============================ 8. SCALABILITY ============================
banner("8. Scalability & Deployment View", NAVY_H)
data_table(["#", "Concern", "Prototype (today)", "Scale-out path"],
    [[str(i), c1, c2, c3] for i, (c1, c2, c3) in enumerate([
        ("API", "Single uvicorn on :8010", "N replicas behind LB (stateless)"),
        ("DB", "SQLite file", "Postgres/Aurora (swap database_url)"),
        ("Pipeline", "In-process sequential", "Each stage \u2192 queue consumer (Celery/SQS/Kafka)"),
        ("Blob storage", "data/blobs/ local", "S3 / Azure Blob (blob_dir abstraction)"),
        ("LLM", "Stub / agent-file (no key)", "OpenAI / Azure OpenAI structured outputs"),
        ("Extraction", "PyMuPDF", "Azure Document Intelligence for scanned PDFs"),
        ("Market data", "Offline scripts \u2192 JSON \u2192 DB", "Scheduled workers + caching layer"),
    ], 1)],
    widths=[0.4, 1.4, 2.4, 2.6], header_hex=NAVY_H, bold_cols=(1,))
para("Pipeline stages are already discrete idempotent functions, so promoting them to independent queue "
     "consumers is a mechanical change, not a redesign.", italic=True, color=AMBER)

doc.add_page_break()

# ============================ 9. DECISIONS & RISKS ============================
banner("9.1 Design Decisions & Trade-offs", NAVY_H)
data_table(["#", "Decision", "Rationale"],
    [[str(i), d, rr] for i, (d, rr) in enumerate([
        ("SQLite first", "Fastest path to a working, queryable memory; ORM keeps Postgres one line away."),
        ("Stub/agent LLM providers", "Entire pipeline runs with NO API key \u2192 demoable, privacy-preserving; real providers plug in."),
        ("period_ordinal integer clock", "Simplifies temporal joins & guidance due-dating vs dates."),
        ("Offline ranking scripts", "Market data fetched in controlled batches and persisted; request path stays fast & resilient."),
        ("Sustainable CAGR over raw YoY for PEG", "Avoids false 'undervalued' signals from one-off spikes (NEOGEN 375%, SAMBHV 162%)."),
    ], 1)],
    widths=[0.4, 2.3, 4.1], header_hex=NAVY_H, bold_cols=(1,))

banner("9.2 Risks & Mitigations", AMBER_H, size=12)
data_table(["#", "Risk", "Mitigation"],
    [[str(i), rk, mg] for i, (rk, mg) in enumerate([
        ("Transcript unavailable (recent IPO / only notices)", "Mark company blocked with reason (HBLENGINE, E2E)."),
        ("Third-party rate-limits / login-gates", "Offline batch fetch + persist; degrade to N/A, never fabricate."),
        ("LLM hallucination", "Structured-output schema + evidence quotes stored with every fact."),
        ("Inflated CAGR for <3y listings", "Flag *, winsorize at 100% in composite."),
        ("Data drift between quarters", "Idempotent pipeline re-run + as-of scoring snapshots."),
    ], 1)],
    widths=[0.4, 3.0, 3.4], header_hex=AMBER_H, bold_cols=(1,))

banner("9.3 Future Enhancements", GREEN_H, size=12)
data_table(["#", "Enhancement"],
    [[str(i), f] for i, f in enumerate([
        "Promote pipeline stages to a real message queue (event-per-stage).",
        "Postgres migration + Alembic migrations (replace ad-hoc ALTER TABLE).",
        "Sortable/filterable ranking columns; sector-relative valuation percentiles.",
        "Real Screener 5-yr Sales/Profit CAGR + ROE ingestion.",
        "AuthN/AuthZ, multi-user watchlists, and audit trail.",
    ], 1)],
    widths=[0.4, 6.4], header_hex=GREEN_H, bold_cols=())

para("")
para("Architecture Design Document \u2014 Concall Intelligence Platform. "
     "Prepared by GitHub Copilot CLI for @vchaurasia_microsoft.", italic=True, color=GREY, size=9)

doc.save(OUT)
print("SAVED", OUT)
